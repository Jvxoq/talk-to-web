"""Unit tests for the file-type text extractors."""

from pathlib import Path

import pytest
from docx import Document as DocxDocument

from app.adapters.extraction.composite_extractor import CompositeTextExtractor
from app.adapters.extraction.docx_extractor import DocxTextExtractor
from app.adapters.extraction.plain_text_extractor import PlainTextExtractor
from app.domain.ingestion.errors import UnsupportedDocumentType
from tests.fakes import FakeTextExtractor


class TestPlainTextExtractor:
    async def test_extracts_txt_content(self, tmp_path: Path) -> None:
        target = tmp_path / "notes.txt"
        target.write_text("hello world", encoding="utf-8")

        text = await PlainTextExtractor().extract(str(target))

        assert text == "hello world"

    async def test_extracts_markdown_as_is(self, tmp_path: Path) -> None:
        target = tmp_path / "notes.md"
        target.write_text("# heading\n\nsome *body* text", encoding="utf-8")

        text = await PlainTextExtractor().extract(str(target))

        assert text == "# heading\n\nsome *body* text"


class TestDocxTextExtractor:
    async def test_extracts_paragraph_text(self, tmp_path: Path) -> None:
        target = tmp_path / "report.docx"
        document = DocxDocument()
        document.add_paragraph("first paragraph")
        document.add_paragraph("second paragraph")
        document.save(str(target))

        text = await DocxTextExtractor().extract(str(target))

        assert text == "first paragraph\nsecond paragraph"


class TestCompositeTextExtractor:
    async def test_dispatches_by_suffix(self) -> None:
        pdf_extractor = FakeTextExtractor("pdf text")
        txt_extractor = FakeTextExtractor("txt text")
        composite = CompositeTextExtractor({".pdf": pdf_extractor, ".txt": txt_extractor})

        assert await composite.extract("uploads/1/report.pdf") == "pdf text"
        assert await composite.extract("uploads/1/notes.txt") == "txt text"

    async def test_dispatch_is_case_insensitive(self) -> None:
        extractor = FakeTextExtractor("text")
        composite = CompositeTextExtractor({".pdf": extractor})

        assert await composite.extract("uploads/1/REPORT.PDF") == "text"

    async def test_raises_for_an_unregistered_suffix(self) -> None:
        composite = CompositeTextExtractor({".pdf": FakeTextExtractor("text")})

        with pytest.raises(UnsupportedDocumentType):
            await composite.extract("uploads/1/archive.zip")
